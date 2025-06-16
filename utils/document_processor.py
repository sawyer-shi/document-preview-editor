#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Document Preview Editor
Copyright (c) 2025 sawyer-shi
Licensed under the Apache License, Version 2.0
https://github.com/sawyer-shi/document-preview-editor

文档处理器模块 - 增强版
Enhanced Document Processor Module

支持的文件格式：
- DOCX: 完全支持，推荐使用
- TXT: 完全支持，自动转换为DOCX格式
- DOC: 有限支持，可能出现格式问题，建议转换为DOCX格式后使用

注意：DOC格式是微软的专有二进制格式，完全解析需要复杂的逆向工程。
为获得最佳体验，建议将DOC文件在Microsoft Word中另存为DOCX格式后使用。

Copyright 2025 sawyer-shi
Licensed under the Apache License, Version 2.0
https://github.com/sawyer-shi
"""

import os
import io
import base64
import tempfile
import zipfile
import subprocess
import shutil
from typing import List, Dict, Any, Tuple, Optional
from docx import Document
from docx.shared import RGBColor, Inches, Pt
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml
import xml.etree.ElementTree as ET
from PIL import Image
from utils.i18n import get_text
import docx2txt

# 条件导入 Windows 特有的模块
try:
    if os.name == 'nt':  # 只在 Windows 环境下导入
        import win32com.client as win32
        import pythoncom
        WINDOWS_COM_AVAILABLE = True
    else:
        WINDOWS_COM_AVAILABLE = False
except ImportError:
    WINDOWS_COM_AVAILABLE = False

class EnhancedWordProcessor:
    """增强的Word文档处理器"""
    
    def __init__(self):
        self.original_doc = None
        self.modified_doc = None
        self.modifications = []
        self.images = {}  # 存储文档中的图片
        self.tables = []  # 存储表格数据
        self.styles = {}  # 存储样式信息
    
    def load_document(self, file_path: str) -> Tuple[bool, str]:
        """加载文档，支持.docx和.txt格式"""
        try:
            # 检查文件格式
            file_ext = os.path.splitext(file_path)[1].lower()
            processed_file_path = file_path
            
            # 检查是否为支持的格式
            if file_ext not in ['.docx', '.txt']:
                return False, f"{get_text('upload_failed')}: 不支持的文件格式。{get_text('invalid_file_format')}"
            
            # 如果是.txt格式，转换为.docx
            if file_ext == '.txt':
                processed_file_path = self._convert_txt_to_docx(file_path)
                if not processed_file_path:
                    return False, f"{get_text('upload_failed')}: 无法处理.txt文件"
            
            # 加载文档
            self.original_doc = Document(processed_file_path)
            
            # 提取文档中的图片（仅对docx格式）
            if file_ext in ['.docx', '.txt']:  # txt转换后也是docx格式
                self._extract_images(processed_file_path)
            
            # 提取样式信息
            self._extract_styles()
            
            return True, get_text('upload_success')
            
        except Exception as e:
            return False, f"{get_text('upload_failed')}: {str(e)}"
    
    def _convert_doc_to_docx(self, doc_path: str) -> Optional[str]:
        """将.doc文件转换为.docx文件 - 改进版"""
        try:
            from utils.i18n import get_text
            print(f"{get_text('doc_conversion_notice')}")
            print(f"DOC文件路径: {doc_path}")
            
            # 首先显示格式警告
            print(f"⚠️  {get_text('doc_format_warning')}")
            
            # 方法1: 使用LibreOffice命令行工具（最可靠）
            libreoffice_result = self._convert_with_libreoffice(doc_path)
            if libreoffice_result:
                print(f"✅ {get_text('doc_conversion_limited')}")
                return libreoffice_result
            
            # 方法2: 使用win32com（仅Windows，但很可靠）
            if WINDOWS_COM_AVAILABLE:
                win32_result = self._convert_with_win32com(doc_path)
                if win32_result:
                    print(f"✅ {get_text('doc_conversion_limited')}")
                    return win32_result
            
            # 方法3: Python库转换（改进版）
            python_result = self._convert_with_python_libraries(doc_path)
            if python_result:
                print(f"✅ {get_text('doc_conversion_limited')}")
                return python_result
            
            # 方法4: 尝试其他方法
            alternative_result = self._convert_with_alternative_tools(doc_path)
            if alternative_result:
                print(f"✅ {get_text('doc_conversion_limited')}")
                return alternative_result
            
            # 如果所有方法都失败，提供用户友好的错误信息
            print("❌ DOC文件转换失败")
            print("💡 建议解决方案：")
            print("   1. 在Microsoft Word中打开DOC文件")
            print("   2. 选择 '文件' -> '另存为'")
            print("   3. 选择 'Word文档(*.docx)' 格式")
            print("   4. 保存后使用DOCX文件")
            
            return None
            
        except Exception as e:
            print(f"DOC转换过程出错: {str(e)}")
            return None
    
    def _convert_txt_to_docx(self, txt_path: str) -> Optional[str]:
        """将.txt文件转换为.docx文件"""
        try:
            print(f"转换TXT文件为DOCX: {txt_path}")
            
            # 读取txt文件内容，尝试多种编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'utf-8-sig', 'cp1252', 'latin1']
            content = None
            used_encoding = None
            
            for encoding in encodings:
                try:
                    with open(txt_path, 'r', encoding=encoding) as f:
                        content = f.read()
                        used_encoding = encoding
                        print(f"成功使用编码 {encoding} 读取文件")
                        break
                except UnicodeDecodeError:
                    continue
                except Exception as e:
                    print(f"使用编码 {encoding} 读取文件失败: {str(e)}")
                    continue
            
            if content is None:
                print("无法读取TXT文件内容，尝试所有编码都失败")
                return None
            
            # 创建新的docx文档
            doc = Document()
            
            # 设置默认样式
            style = doc.styles['Normal']
            font = style.font
            font.name = '宋体'  # 中文友好字体
            font.size = Pt(12)
            
            # 分割文本为段落（按换行符分割）
            paragraphs = content.split('\n')
            
            for para_text in paragraphs:
                # 去除行首行尾空白，但保留段落内部的空格
                para_text = para_text.rstrip()
                
                # 添加段落到文档
                paragraph = doc.add_paragraph()
                
                # 如果段落为空，添加一个空行
                if not para_text.strip():
                    continue
                
                # 添加文本运行
                run = paragraph.add_run(para_text)
                
                # 设置字体属性以确保中文显示正常
                run.font.name = '宋体'
                run.font.size = Pt(12)
                
                # 设置段落格式
                paragraph_format = paragraph.paragraph_format
                paragraph_format.line_spacing = 1.15  # 行间距
                paragraph_format.space_after = Pt(6)  # 段后间距
            
            # 如果文档为空，至少添加一个空段落
            if len(doc.paragraphs) == 0:
                doc.add_paragraph()
            
            # 生成输出文件路径
            temp_dir = tempfile.gettempdir()
            base_name = os.path.splitext(os.path.basename(txt_path))[0]
            output_path = os.path.join(temp_dir, f'{base_name}_converted.docx')
            
            # 保存转换后的文档
            doc.save(output_path)
            
            print(f"TXT转DOCX成功: {output_path}")
            print(f"使用编码: {used_encoding}, 段落数: {len(doc.paragraphs)}")
            
            return output_path
            
        except Exception as e:
            print(f"TXT转DOCX失败: {str(e)}")
            return None
    
    def _is_valid_doc_file(self, file_path: str) -> bool:
        """检查是否为有效的DOC文件 - 改进版本"""
        try:
            print(f"验证DOC文件格式: {file_path}")
            
            with open(file_path, 'rb') as f:
                # 读取更多的文件头信息
                f.seek(0)
                header = f.read(512)  # 读取512字节用于全面检查
                
                # DOC文件的各种签名
                doc_signatures = [
                    b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1',  # OLE2 signature (最常见)
                    b'\x0d\x44\x4f\x43',  # DOC signature
                    b'\xdb\xa5\x2d\x00',  # DOC variant
                    b'\xec\xa5',  # Old DOC format
                    b'\xfe\x37',  # Old format variant
                    b'\x31\xbe\x00\x00'   # Another DOC variant
                ]
                
                # 1. 检查OLE2格式（最常见的DOC格式）
                if header.startswith(b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1'):
                    print("检测到OLE2格式DOC文件")
                    return True
                
                # 2. 检查其他DOC签名
                for i, sig in enumerate(doc_signatures):
                    if header.startswith(sig):
                        print(f"检测到DOC签名 #{i+1}")
                        return True
                
                # 3. 检查RTF格式（有些.doc实际是RTF）
                if header.startswith(b'{\\rtf') or b'{\\rtf' in header[:100]:
                    print("检测到RTF格式文件（扩展名为.doc）")
                    return True
                
                # 4. 检查是否为Word 97-2003格式的其他变体
                # 某些DOC文件可能有不同的头部，但包含Word特有的标识
                word_indicators = [
                    b'Microsoft Office Word',
                    b'Word.Document',
                    b'MSWordDoc',
                    b'\\x00W\\x00o\\x00r\\x00d',  # Unicode "Word"
                    b'PK\\x03\\x04'  # 这实际上是ZIP文件，可能是错误命名的DOCX
                ]
                
                for indicator in word_indicators:
                    if indicator in header:
                        if indicator == b'PK\\x03\\x04':
                            print("文件似乎是DOCX格式，但扩展名为.doc")
                        else:
                            print(f"检测到Word文档标识: {indicator}")
                        return True
                
                # 5. 宽松验证：如果是.doc扩展名且不是明显的其他格式
                if file_path.lower().endswith('.doc'):
                    # 检查是否为其他已知格式
                    known_formats = [
                        (b'%PDF', 'PDF'),
                        (b'\\x89PNG', 'PNG'),
                        (b'\\xff\\xd8\\xff', 'JPEG'),
                        (b'GIF8', 'GIF'),
                        (b'BM', 'BMP'),
                        (b'\\x1f\\x8b', 'GZIP'),
                        (b'\\x50\\x4b\\x03\\x04', 'ZIP/DOCX')
                    ]
                    
                    for sig, format_name in known_formats:
                        if header.startswith(sig.encode() if isinstance(sig, str) else sig):
                            if format_name == 'ZIP/DOCX':
                                print("文件实际上是DOCX格式，将作为DOCX处理")
                                return True  # 允许处理，但会在后续转换中处理
                            else:
                                print(f"文件实际上是{format_name}格式，不是DOC")
                                return False
                    
                    # 如果没有检测到其他格式，就认为可能是DOC
                    print("未检测到明确的DOC签名，但扩展名正确，尝试作为DOC处理")
                    return True
                
                print("文件不是有效的DOC格式")
                return False
                
        except Exception as e:
            print(f"检查DOC文件时出错: {str(e)}")
            # 如果检查过程出错，但扩展名正确，就尝试处理
            if file_path.lower().endswith('.doc'):
                print("DOC文件检查异常，但扩展名正确，允许尝试转换")
                return True
            return False
    
    def _convert_with_libreoffice(self, doc_path: str) -> Optional[str]:
        """使用LibreOffice转换DOC文件"""
        try:
            # 检查LibreOffice可执行文件
            libreoffice_paths = [
                'libreoffice',
                'soffice',
                '/usr/bin/libreoffice',
                '/usr/bin/soffice',
                '/Applications/LibreOffice.app/Contents/MacOS/soffice',
                'C:\\Program Files\\LibreOffice\\program\\soffice.exe',
                'C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe'
            ]
            
            libreoffice_cmd = None
            for path in libreoffice_paths:
                if shutil.which(path) or os.path.exists(path):
                    libreoffice_cmd = path
                    break
            
            if not libreoffice_cmd:
                return None
            
            # 执行转换
            temp_dir = os.path.dirname(doc_path)
            cmd = [
                libreoffice_cmd,
                '--headless',
                '--convert-to', 'docx',
                '--outdir', temp_dir,
                doc_path
            ]
            
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
            
            # 检查转换结果
            base_name = os.path.splitext(os.path.basename(doc_path))[0]
            converted_path = os.path.join(temp_dir, f"{base_name}.docx")
            
            if os.path.exists(converted_path) and os.path.getsize(converted_path) > 0:
                return converted_path
                
            return None
            
        except Exception as e:
            print(f"LibreOffice转换失败: {str(e)}")
            return None
    
    def _convert_with_win32com(self, doc_path: str) -> Optional[str]:
        """使用Windows COM组件转换DOC文件"""
        try:
            # 检查 Windows COM 是否可用
            if not WINDOWS_COM_AVAILABLE:
                print("win32com不可用 - 不是Windows环境或未安装")
                return None
            
            # 创建Word应用程序对象
            word_app = win32.Dispatch("Word.Application")
            word_app.Visible = False
            
            # 打开DOC文件
            doc = word_app.Documents.Open(os.path.abspath(doc_path))
            
            # 转换为DOCX格式
            temp_docx_path = doc_path.replace('.doc', '_converted.docx')
            doc.SaveAs2(os.path.abspath(temp_docx_path), FileFormat=16)  # 16 = docx格式
            
            # 关闭文档和应用程序
            doc.Close()
            word_app.Quit()
            
            if os.path.exists(temp_docx_path):
                return temp_docx_path
                
            return None
            
        except ImportError:
            print("win32com不可用")
            return None
        except Exception as e:
            print(f"win32com转换失败: {str(e)}")
            return None
    
    def _convert_with_python_libraries(self, doc_path: str) -> Optional[str]:
        """使用Python库转换DOC文件 - 改进版"""
        try:
            # 方法1: 尝试使用antiword（如果可用）
            antiword_result = self._try_antiword_conversion(doc_path)
            if antiword_result:
                return antiword_result
            
            # 方法2: 尝试使用catdoc（如果可用）
            catdoc_result = self._try_catdoc_conversion(doc_path)
            if catdoc_result:
                return catdoc_result
            
            # 方法3: 使用docx2txt（适用于某些DOC文件）
            docx2txt_result = self._try_docx2txt_conversion(doc_path)
            if docx2txt_result:
                return docx2txt_result
            
            # 方法4: 使用oletools/olefile读取OLE2格式（改进版）
            ole_result = self._try_improved_ole_conversion(doc_path)
            if ole_result:
                return ole_result
            
            # 方法5: 智能二进制分析提取文本（最后的备选方案）
            binary_result = self._try_smart_binary_extraction(doc_path)
            if binary_result:
                return binary_result
            
            return None
            
        except Exception as e:
            print(f"Python库转换失败: {str(e)}")
            return None
    
    def _try_antiword_conversion(self, doc_path: str) -> Optional[str]:
        """尝试使用antiword转换DOC文件"""
        try:
            import subprocess
            import tempfile
            
            # 创建临时文件存储转换结果
            with tempfile.NamedTemporaryFile(mode='w+', suffix='.txt', delete=False) as temp_file:
                temp_path = temp_file.name
            
            try:
                # 尝试调用antiword命令
                result = subprocess.run(
                    ['antiword', doc_path, '-t', temp_path],
                    capture_output=True,
                    text=True,
                    timeout=30
                )
                
                if result.returncode == 0:
                    with open(temp_path, 'r', encoding='utf-8') as f:
                        text_content = f.read()
                    
                    if text_content.strip():
                        print("antiword转换成功")
                        return self._create_docx_from_text(text_content, doc_path, "antiword转换")
                
            except (subprocess.TimeoutExpired, FileNotFoundError):
                pass  # antiword不可用或超时
            finally:
                # 清理临时文件
                try:
                    os.unlink(temp_path)
                except:
                    pass
                    
        except Exception as e:
            print(f"antiword转换失败: {str(e)}")
        
        return None
    
    def _try_catdoc_conversion(self, doc_path: str) -> Optional[str]:
        """尝试使用catdoc转换DOC文件"""
        try:
            import subprocess
            
            try:
                # 尝试调用catdoc命令
                result = subprocess.run(
                    ['catdoc', doc_path],
                    capture_output=True,
                    text=True,
                    timeout=30
                )
                
                if result.returncode == 0 and result.stdout.strip():
                    text_content = result.stdout
                    print("catdoc转换成功")
                    return self._create_docx_from_text(text_content, doc_path, "catdoc转换")
                
            except (subprocess.TimeoutExpired, FileNotFoundError):
                pass  # catdoc不可用或超时
                    
        except Exception as e:
            print(f"catdoc转换失败: {str(e)}")
        
        return None
    
    def _try_docx2txt_conversion(self, doc_path: str) -> Optional[str]:
        """尝试使用docx2txt转换"""
        try:
            import docx2txt
            
            # 有些DOC文件实际上是ZIP格式，可以用docx2txt处理
            text_content = docx2txt.process(doc_path)
            
            if text_content and text_content.strip():
                return self._create_docx_from_text(text_content, doc_path, "docx2txt转换")
                
        except Exception as e:
            if "not a zip file" not in str(e).lower():
                print(f"docx2txt尝试失败: {str(e)}")
        
        return None
    
    def _try_ole_conversion(self, doc_path: str) -> Optional[str]:
        """尝试使用OLE工具转换"""
        try:
            # 尝试使用olefile库读取OLE2文档
            try:
                import olefile
                
                if olefile.isOleFile(doc_path):
                    print("检测到OLE2文件，尝试提取内容")
                    
                    # 打开OLE文件
                    ole = olefile.OleFileIO(doc_path)
                    
                    # 查找Word文档流
                    streams = ole.listdir()
                    print(f"OLE流列表: {streams}")
                    
                    # 寻找主要的Word文档流
                    doc_streams = ['WordDocument', '1Table', '0Table', 'Data']
                    text_content = ""
                    
                    for stream_name in doc_streams:
                        try:
                            if ole._olestream_size(stream_name) > 0:
                                stream_data = ole._olestream(stream_name).read()
                                # 简单的文本提取（这里可以改进）
                                text_part = self._extract_text_from_stream(stream_data)
                                if text_part:
                                    text_content += text_part + "\n"
                        except:
                            continue
                    
                    ole.close()
                    
                    if text_content.strip():
                        return self._create_docx_from_text(text_content, doc_path, "OLE提取")
                
            except ImportError:
                print("olefile库不可用")
                
        except Exception as e:
            print(f"OLE转换失败: {str(e)}")
        
        return None
    
    def _try_improved_ole_conversion(self, doc_path: str) -> Optional[str]:
        """改进的OLE文件转换方法"""
        try:
            # 尝试使用olefile库读取OLE2文档
            try:
                import olefile
                
                if olefile.isOleFile(doc_path):
                    print("检测到OLE2文件，使用改进的提取方法")
                    
                    # 打开OLE文件
                    ole = olefile.OleFileIO(doc_path)
                    
                    # 查找Word文档流
                    streams = ole.listdir()
                    print(f"OLE流列表: {streams}")
                    
                    text_content = ""
                    
                    # 优先查找WordDocument流
                    if ole._olestream_size('WordDocument') > 0:
                        word_doc_stream = ole._olestream('WordDocument').read()
                        
                        # Word文档的文本通常从特定偏移开始
                        # 跳过文档头部信息，查找文本内容
                        text_content = self._extract_word_document_text(word_doc_stream)
                    
                    # 如果WordDocument流没有找到文本，尝试其他流
                    if not text_content.strip():
                        for stream_name in ['1Table', '0Table', 'Data']:
                            try:
                                if ole._olestream_size(stream_name) > 0:
                                    stream_data = ole._olestream(stream_name).read()
                                    text_part = self._extract_text_from_stream(stream_data)
                                    if text_part and len(text_part.strip()) > 10:
                                        text_content += text_part + "\n"
                            except:
                                continue
                    
                    ole.close()
                    
                    if text_content.strip():
                        print("改进的OLE提取成功")
                        return self._create_docx_from_text(text_content, doc_path, "改进OLE提取")
                
            except ImportError:
                print("olefile库不可用")
                
        except Exception as e:
            print(f"改进OLE转换失败: {str(e)}")
        
        return None
    
    def _extract_word_document_text(self, word_doc_data: bytes) -> str:
        """从WordDocument流中提取文本"""
        try:
            text_content = ""
            
            # Word文档的文本通常存储在特定的结构中
            # 这是一个简化的实现，实际的Word格式非常复杂
            
            # 方法1: 查找Unicode文本
            import re
            
            # 查找UTF-16编码的文本
            try:
                # 尝试UTF-16LE解码
                decoded = word_doc_data.decode('utf-16le', errors='ignore')
                # 清理控制字符，保留中文和英文
                cleaned = re.sub(r'[^\u4e00-\u9fff\u0020-\u007e\n\r\t，。！？；：""''（）【】《》]', '', decoded)
                
                # 提取有意义的文本段
                lines = cleaned.split('\n')
                meaningful_lines = []
                
                for line in lines:
                    line = line.strip()
                    if len(line) >= 2:  # 至少2个字符
                        # 检查是否包含中文或有意义的英文
                        if (re.search(r'[\u4e00-\u9fff]', line) or  # 包含中文
                            (re.search(r'[a-zA-Z]{2,}', line) and len(line) >= 3)):  # 包含英文单词
                            meaningful_lines.append(line)
                
                if meaningful_lines:
                    text_content = '\n'.join(meaningful_lines)
                    
            except:
                pass
            
            # 方法2: 如果UTF-16失败，尝试其他编码
            if not text_content.strip():
                encodings = ['utf-8', 'gbk', 'gb2312', 'big5']
                
                for encoding in encodings:
                    try:
                        decoded = word_doc_data.decode(encoding, errors='ignore')
                        cleaned = re.sub(r'[^\u4e00-\u9fff\u0020-\u007e\n\r\t，。！？；：""''（）【】《》]', '', decoded)
                        
                        if len(cleaned.strip()) > len(text_content.strip()):
                            text_content = cleaned
                            
                    except:
                        continue
            
            return text_content.strip()
            
        except Exception as e:
            print(f"WordDocument文本提取失败: {str(e)}")
            return ""
    
    def _try_smart_binary_extraction(self, doc_path: str) -> Optional[str]:
        """智能二进制文本提取 - 最后的备选方案"""
        try:
            with open(doc_path, 'rb') as f:
                data = f.read()
            
            print(f"开始智能二进制文本提取，文件大小: {len(data)} 字节")
            
            import re
            meaningful_texts = []
            
            # 方法1: 查找UTF-16编码的文本（Word常用）
            try:
                # 分段查找UTF-16文本
                segment_size = 1024
                for i in range(0, len(data) - 1, segment_size):
                    segment = data[i:i+segment_size]
                    
                    # 尝试UTF-16LE解码
                    try:
                        decoded = segment.decode('utf-16le', errors='ignore')
                        # 查找有意义的文本
                        lines = decoded.split('\x00')  # UTF-16中的空字符
                        for line in lines:
                            line = line.strip()
                            if len(line) >= 2:
                                # 检查是否包含中文或有意义的英文
                                if (re.search(r'[\u4e00-\u9fff]', line) or
                                    (re.search(r'[a-zA-Z]{3,}', line) and 
                                     not re.match(r'^[A-Z]+$', line) and  # 不是全大写
                                     'Microsoft' not in line and 'Office' not in line)):  # 不是软件标识
                                    meaningful_texts.append(line)
                    except:
                        continue
                        
            except Exception as e:
                print(f"UTF-16提取失败: {str(e)}")
            
            # 方法2: 查找UTF-8编码的中文文本
            try:
                chinese_pattern = re.compile(rb'[\xe4-\xe9][\x80-\xbf]{2}+')
                chinese_matches = chinese_pattern.findall(data)
                
                for match in chinese_matches:
                    try:
                        decoded = match.decode('utf-8', errors='ignore').strip()
                        if len(decoded) >= 2 and '\u4e00' <= decoded[0] <= '\u9fff':
                            meaningful_texts.append(decoded)
                    except:
                        continue
                        
            except Exception as e:
                print(f"中文文本提取失败: {str(e)}")
            
            # 方法3: 查找GBK编码的中文文本
            try:
                # GBK编码的中文字符范围
                gbk_pattern = re.compile(rb'[\xa1-\xfe][\xa1-\xfe]+')
                gbk_matches = gbk_pattern.findall(data)
                
                for match in gbk_matches:
                    try:
                        decoded = match.decode('gbk', errors='ignore').strip()
                        if len(decoded) >= 2 and '\u4e00' <= decoded[0] <= '\u9fff':
                            meaningful_texts.append(decoded)
                    except:
                        continue
                        
            except Exception as e:
                print(f"GBK文本提取失败: {str(e)}")
            
            # 方法4: 查找ASCII文本（过滤格式标记）
            try:
                ascii_pattern = re.compile(rb'[a-zA-Z][a-zA-Z0-9\s\.,;:\-]{2,30}[a-zA-Z0-9]')
                ascii_matches = ascii_pattern.findall(data)
                
                for match in ascii_matches:
                    try:
                        decoded = match.decode('ascii', errors='ignore').strip()
                        # 过滤掉明显的格式标记
                        if (len(decoded) >= 3 and 
                            not re.match(r'^[A-Z]{2,}$', decoded) and  # 不是全大写缩写
                            not re.match(r'^[0-9\.,\-\s]+$', decoded) and  # 不是纯数字
                            'Microsoft' not in decoded and 'Office' not in decoded and
                            'Word' not in decoded and 'Document' not in decoded and
                            'Normal' not in decoded and 'Table' not in decoded):
                            meaningful_texts.append(decoded)
                    except:
                        continue
                        
            except Exception as e:
                print(f"ASCII文本提取失败: {str(e)}")
            
            # 去重并组合文本
            if meaningful_texts:
                # 去重
                unique_texts = list(dict.fromkeys(meaningful_texts))
                
                # 过滤掉太短的文本
                filtered_texts = [text for text in unique_texts if len(text.strip()) >= 2]
                
                if filtered_texts:
                    # 组合文本
                    combined_text = '\n'.join(filtered_texts)
                    
                    print(f"智能提取完成，提取到 {len(filtered_texts)} 个文本段")
                    print(f"文本预览: {combined_text[:100]}...")
                    
                    return self._create_docx_from_text(combined_text, doc_path, "智能二进制提取")
            
            print("智能二进制提取未找到有效文本")
            return None
                
        except Exception as e:
            print(f"二进制文本提取失败: {str(e)}")
            return None
    
    def _convert_with_alternative_tools(self, doc_path: str) -> Optional[str]:
        """使用替代工具转换DOC文件"""
        try:
            # 尝试使用命令行工具
            tools = [
                ('antiword', ['-t', doc_path]),  # antiword工具
                ('catdoc', ['-w', doc_path]),    # catdoc工具
                ('wvText', [doc_path, '-']),     # wvWare工具
            ]
            
            for tool_name, cmd_args in tools:
                if shutil.which(tool_name):
                    try:
                        result = subprocess.run(
                            cmd_args, 
                            capture_output=True, 
                            text=True, 
                            timeout=30,
                            encoding='utf-8',
                            errors='ignore'
                        )
                        
                        if result.returncode == 0 and result.stdout.strip():
                            text_content = result.stdout.strip()
                            if len(text_content) > 20:
                                print(f"使用{tool_name}工具转换成功")
                                return self._create_docx_from_text(text_content, doc_path, f"{tool_name}转换")
                    except Exception as e:
                        print(f"{tool_name}工具转换失败: {str(e)}")
                        continue
            
            return None
            
        except Exception as e:
            print(f"替代工具转换失败: {str(e)}")
            return None
    
    def _extract_text_from_stream(self, stream_data: bytes) -> str:
        """从OLE流中提取文本"""
        try:
            # 简单的文本提取方法
            text = ""
            
            # 尝试直接解码
            try:
                text = stream_data.decode('utf-8', errors='ignore')
            except:
                try:
                    text = stream_data.decode('latin-1', errors='ignore')
                except:
                    pass
            
            # 清理文本
            import re
            text = re.sub(r'[^\x20-\x7E\n\r\t]', '', text)
            text = re.sub(r'\s+', ' ', text).strip()
            
            return text if len(text) > 10 else ""
            
        except Exception as e:
            print(f"从流中提取文本失败: {str(e)}")
            return ""
    
    def _create_docx_from_text(self, text_content: str, original_path: str, method_name: str) -> str:
        """从文本内容创建DOCX文档"""
        try:
            # 创建新的DOCX文档
            new_doc = Document()
            
            # 添加文档标题
            title_para = new_doc.add_paragraph()
            title_run = title_para.add_run(f"从DOC文件转换（{method_name}）")
            title_run.bold = True
            title_para.alignment = 1  # 居中对齐
            
            # 添加分隔线
            new_doc.add_paragraph("=" * 50)
            new_doc.add_paragraph()  # 空行
            
            # 处理文本内容，分段落
            paragraphs = text_content.split('\n')
            current_para_text = ""
            
            for line in paragraphs:
                line = line.strip()
                if line:
                    if current_para_text:
                        current_para_text += " " + line
                    else:
                        current_para_text = line
                    
                    # 根据长度或标点创建段落
                    if (len(current_para_text) > 300 or 
                        line.endswith('.') or line.endswith('。') or 
                        line.endswith('!') or line.endswith('！') or
                        line.endswith('?') or line.endswith('？')):
                        
                        new_doc.add_paragraph(current_para_text)
                        current_para_text = ""
                else:
                    # 空行，结束当前段落
                    if current_para_text:
                        new_doc.add_paragraph(current_para_text)
                        current_para_text = ""
            
            # 添加剩余内容
            if current_para_text:
                new_doc.add_paragraph(current_para_text)
            
            # 保存转换后的文档
            temp_docx_path = original_path.replace('.doc', f'_converted_{method_name}.docx')
            new_doc.save(temp_docx_path)
            
            print(f"使用{method_name}创建DOCX文档: {temp_docx_path}")
            return temp_docx_path
            
        except Exception as e:
            print(f"创建DOCX文档失败: {str(e)}")
            return None
    
    def _extract_images(self, docx_path: str):
        """提取文档中的图片"""
        try:
            with zipfile.ZipFile(docx_path, 'r') as docx_zip:
                # 查找media文件夹中的图片
                media_files = [f for f in docx_zip.namelist() if f.startswith('word/media/')]
                
                for media_file in media_files:
                    try:
                        # 读取图片数据
                        img_data = docx_zip.read(media_file)
                        
                        # 转换为base64编码
                        img_base64 = base64.b64encode(img_data).decode('utf-8')
                        
                        # 确定图片类型
                        img_ext = os.path.splitext(media_file)[1].lower()
                        if img_ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp']:
                            mime_type = f"image/{img_ext[1:]}"
                            if img_ext == '.jpg':
                                mime_type = "image/jpeg"
                            
                            # 存储图片信息
                            filename = os.path.basename(media_file)
                            self.images[filename] = {
                                'data': img_base64,
                                'mime_type': mime_type,
                                'filename': filename
                            }
                    except Exception as e:
                        print(f"提取图片失败 {media_file}: {str(e)}")
                        
        except Exception as e:
            print(f"提取图片时出错: {str(e)}")
    
    def _extract_styles(self):
        """提取文档样式信息"""
        try:
            if not self.original_doc:
                return
            
            # 提取段落样式
            for paragraph in self.original_doc.paragraphs:
                if paragraph.style:
                    style_name = paragraph.style.name
                    self.styles[style_name] = {
                        'type': 'paragraph',
                        'name': style_name
                    }
            
            # 提取字符样式
            for paragraph in self.original_doc.paragraphs:
                for run in paragraph.runs:
                    if run.style:
                        style_name = run.style.name
                        self.styles[style_name] = {
                            'type': 'character',
                            'name': style_name
                        }
                        
        except Exception as e:
            print(f"提取样式时出错: {str(e)}")
    
    def extract_content_with_formatting(self, doc: Document) -> List[Dict[str, Any]]:
        """提取文档内容，保持复杂格式，实现Word编辑器样式"""
        content = []
        
        try:
            # 处理段落
            for para_idx, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip() or self._has_images(paragraph):
                    para_data = {
                        'type': 'paragraph',
                        'index': para_idx,
                        'text': paragraph.text,
                        'style': paragraph.style.name if paragraph.style else 'Normal',
                        'alignment': self._get_alignment_style(paragraph.alignment),
                        'paragraph_format': self._get_paragraph_format(paragraph),
                        'runs': [],
                        'images': []
                    }
                    
                    # 处理段落中的runs（保持字体格式）
                    for run in paragraph.runs:
                        run_data = {
                            'text': run.text,
                            'bold': run.bold,
                            'italic': run.italic,
                            'underline': run.underline,
                            'font_name': run.font.name or 'Times New Roman',
                            'font_size': run.font.size.pt if run.font.size else 12,
                            'font_color': self._get_font_color(run),
                            'highlight_color': self._get_highlight_color(run),
                            'subscript': run.font.subscript,
                            'superscript': run.font.superscript,
                        }
                        para_data['runs'].append(run_data)
                    
                    # 检查段落中的图片
                    images = self._extract_paragraph_images(paragraph)
                    if images:
                        para_data['images'] = images
                    
                    content.append(para_data)
            
            # 处理表格
            for table_idx, table in enumerate(doc.tables):
                table_data = {
                    'type': 'table',
                    'index': table_idx,
                    'rows': [],
                    'style': table.style.name if table.style else 'Table Grid',
                    'table_format': self._get_table_format(table)
                }
                
                for row_idx, row in enumerate(table.rows):
                    row_data = []
                    for cell_idx, cell in enumerate(row.cells):
                        cell_data = {
                            'text': cell.text,
                            'paragraphs': [],
                            'cell_format': self._get_cell_format(cell)
                        }
                        
                        # 处理单元格中的段落
                        for para in cell.paragraphs:
                            if para.text.strip() or self._has_images(para):
                                para_data = {
                                    'text': para.text,
                                    'style': para.style.name if para.style else 'Normal',
                                    'alignment': self._get_alignment_style(para.alignment),
                                    'runs': []
                                }
                                
                                for run in para.runs:
                                    run_data = {
                                        'text': run.text,
                                        'bold': run.bold,
                                        'italic': run.italic,
                                        'underline': run.underline,
                                        'font_name': run.font.name or 'Times New Roman',
                                        'font_size': run.font.size.pt if run.font.size else 12,
                                        'font_color': self._get_font_color(run),
                                        'highlight_color': self._get_highlight_color(run),
                                    }
                                    para_data['runs'].append(run_data)
                                
                                cell_data['paragraphs'].append(para_data)
                        
                        row_data.append(cell_data)
                    
                    table_data['rows'].append(row_data)
                
                content.append(table_data)
                
        except Exception as e:
            print(f"提取文档内容时出错: {str(e)}")
        
        return content
    
    def _get_alignment_style(self, alignment) -> str:
        """获取段落对齐方式的CSS样式"""
        alignment_map = {
            0: 'left',      # WD_ALIGN_PARAGRAPH.LEFT
            1: 'center',    # WD_ALIGN_PARAGRAPH.CENTER
            2: 'right',     # WD_ALIGN_PARAGRAPH.RIGHT
            3: 'justify',   # WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        
        if alignment is None:
            return 'left'
        
        try:
            return alignment_map.get(int(alignment), 'left')
        except:
            return 'left'
    
    def _get_paragraph_format(self, paragraph) -> Dict[str, Any]:
        """获取段落格式信息"""
        format_info = {
            'space_before': 0,
            'space_after': 0,
            'line_spacing': 1.0,
            'left_indent': 0,
            'right_indent': 0,
            'first_line_indent': 0
        }
        
        try:
            pf = paragraph.paragraph_format
            
            if pf.space_before:
                format_info['space_before'] = pf.space_before.pt
            if pf.space_after:
                format_info['space_after'] = pf.space_after.pt
            if pf.line_spacing:
                format_info['line_spacing'] = pf.line_spacing
            if pf.left_indent:
                format_info['left_indent'] = pf.left_indent.pt
            if pf.right_indent:
                format_info['right_indent'] = pf.right_indent.pt
            if pf.first_line_indent:
                format_info['first_line_indent'] = pf.first_line_indent.pt
                
        except Exception as e:
            print(f"获取段落格式时出错: {str(e)}")
        
        return format_info
    
    def _get_highlight_color(self, run) -> Optional[str]:
        """获取高亮颜色"""
        try:
            if hasattr(run.font, 'highlight_color') and run.font.highlight_color:
                # 简化处理，返回黄色高亮
                return '#FFFF00'
            return None
        except:
            return None
    
    def _get_table_format(self, table) -> Dict[str, Any]:
        """获取表格格式信息"""
        format_info = {
            'width': '100%',
            'border_style': 'solid',
            'border_width': '1px',
            'border_color': '#000000'
        }
        
        try:
            # 这里可以添加更多表格格式提取逻辑
            pass
        except Exception as e:
            print(f"获取表格格式时出错: {str(e)}")
        
        return format_info
    
    def _get_cell_format(self, cell) -> Dict[str, Any]:
        """获取单元格格式信息"""
        format_info = {
            'background_color': None,
            'vertical_alignment': 'top',
            'padding': '5px'
        }
        
        try:
            # 这里可以添加更多单元格格式提取逻辑
            pass
        except Exception as e:
            print(f"获取单元格格式时出错: {str(e)}")
        
        return format_info
    
    def _has_images(self, paragraph) -> bool:
        """检查段落是否包含图片"""
        try:
            # 简化的图片检测方法
            paragraph_xml = paragraph._element.xml
            return 'blip' in paragraph_xml or 'pic:pic' in paragraph_xml
        except Exception as e:
            print(f"检查图片时出错: {str(e)}")
            return False
    
    def _extract_paragraph_images(self, paragraph) -> List[Dict[str, Any]]:
        """提取段落中的图片"""
        images = []
        try:
            # 如果段落包含图片，返回所有可用的图片
            if self._has_images(paragraph):
                # 简化处理：返回文档中的所有图片
                for filename, img_data in self.images.items():
                    images.append({
                        'filename': filename,
                        'data': img_data['data'],
                        'mime_type': img_data['mime_type'],
                        'embed_id': f'img_{len(images)}'
                    })
                    # 限制每个段落最多显示3张图片，避免重复
                    if len(images) >= 3:
                        break
                        
        except Exception as e:
            print(f"提取段落图片时出错: {str(e)}")
        
        return images
    
    def _get_font_color(self, run) -> Optional[str]:
        """获取字体颜色"""
        try:
            if run.font.color and run.font.color.rgb:
                rgb = run.font.color.rgb
                return f"rgb({rgb.r}, {rgb.g}, {rgb.b})"
            return None
        except:
            return None
    
    def apply_modifications(self, modifications: List[Dict[str, str]]) -> Tuple[bool, str]:
        """应用修改到文档，完全保持格式和资源"""
        try:
            if not self.original_doc:
                return False, get_text('file_not_found')
            
            from utils.i18n import get_text
            print(get_text('document_applying_modifications'))
            
            # 方法1：尝试使用高级复制方法（完整保持所有内容）
            success = self._advanced_copy_with_modifications(modifications)
            
            if not success:
                print(get_text('doc_extraction_failed'))
                # 方法2：使用标准方法
                self.modified_doc = Document()
                self._copy_styles()
                self._copy_and_modify_content(modifications)
            
            self.modifications = modifications
            print(get_text('modification_applied_complete'))
            return True, get_text('modifications_applied')
            
        except Exception as e:
            print(f"应用修改时出错: {str(e)}")
            import traceback
            traceback.print_exc()
            return False, f"{get_text('processing_failed')}: {str(e)}"
    
    def _advanced_copy_with_modifications(self, modifications: List[Dict[str, str]]) -> bool:
        """高级文档复制方法，完全保持所有内容"""
        try:
            from docx import Document
            import tempfile
            import os
            from copy import deepcopy
            
            # 创建修改映射
            modification_map = {}
            for mod in modifications:
                modification_map[mod['original_text']] = mod['new_text']
            
            # 保存原文档到临时文件
            temp_path = os.path.join(tempfile.gettempdir(), f"temp_original_{id(self)}.docx")
            self.original_doc.save(temp_path)
            
            # 重新加载文档以创建完整副本
            self.modified_doc = Document(temp_path)
            
            from utils.i18n import get_text
            print(get_text('document_copy_created'))
            
            # 在副本上应用文本修改
            modified_paragraphs = 0
            modified_tables = 0
            
            # 修改段落中的文本
            for paragraph in self.modified_doc.paragraphs:
                for original_text, new_text in modification_map.items():
                    if original_text in paragraph.text:
                        # 查找并替换文本，保持格式
                        self._replace_text_in_paragraph(paragraph, original_text, new_text)
                        modified_paragraphs += 1
            
            # 修改表格中的文本
            for table in self.modified_doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for original_text, new_text in modification_map.items():
                                if original_text in paragraph.text:
                                    self._replace_text_in_paragraph(paragraph, original_text, new_text)
                                    modified_tables += 1
            
            print(f"{get_text('text_modification_complete')} - {get_text('paragraph_replacement')}: {modified_paragraphs}, {get_text('table_replacement')}: {modified_tables}")
            
            # 清理临时文件
            try:
                os.remove(temp_path)
            except:
                pass
            
            return True
            
        except Exception as e:
            print(f"高级复制方法失败: {str(e)}")
            return False
    
    def _replace_text_in_paragraph(self, paragraph, original_text: str, new_text: str):
        """在段落中替换文本，保持格式"""
        try:
            # 检查段落是否包含目标文本
            if original_text not in paragraph.text:
                return
            
            # 遍历段落中的所有runs
            for run in paragraph.runs:
                if original_text in run.text:
                    # 替换文本内容
                    run.text = run.text.replace(original_text, new_text)
                    
                    # 添加高亮标记（如果是修改后的文本）
                    try:
                        from docx.enum.text import WD_COLOR_INDEX
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    except:
                        pass
                    
                    from utils.i18n import get_text
                    print(f"{get_text('paragraph_replacement')}: '{original_text}' -> '{new_text}'")
                    return
            
            # 如果单个run中没有完整的文本，可能文本跨越多个runs
            full_text = paragraph.text
            if original_text in full_text:
                new_full_text = full_text.replace(original_text, new_text)
                
                # 清除现有runs的文本
                for run in paragraph.runs:
                    run.text = ""
                
                # 在第一个run中设置新文本
                if paragraph.runs:
                    paragraph.runs[0].text = new_full_text
                    try:
                        from docx.enum.text import WD_COLOR_INDEX
                        paragraph.runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW
                    except:
                        pass
                else:
                    # 如果没有runs，创建一个新的
                    new_run = paragraph.add_run(new_full_text)
                    try:
                        from docx.enum.text import WD_COLOR_INDEX
                        new_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    except:
                        pass
                
                from utils.i18n import get_text
                print(f"{get_text('text_replacement')}: '{original_text}' -> '{new_text}'")
                    
        except Exception as e:
            print(f"替换段落文本时出错: {str(e)}")
    
    def _copy_styles(self):
        """复制原文档的样式到新文档"""
        try:
            # 复制段落样式
            for style in self.original_doc.styles:
                if style.name not in [s.name for s in self.modified_doc.styles]:
                    try:
                        new_style = self.modified_doc.styles.add_style(style.name, style.type)
                        if hasattr(style, 'font'):
                            new_style.font.name = style.font.name
                            new_style.font.size = style.font.size
                            new_style.font.bold = style.font.bold
                            new_style.font.italic = style.font.italic
                    except:
                        pass  # 忽略样式复制错误
        except Exception as e:
            print(f"复制样式时出错: {str(e)}")
    
    def _copy_and_modify_content(self, modifications: List[Dict[str, str]]):
        """复制并修改文档内容，完全保持所有格式、图片和对象"""
        try:
            # 创建修改映射
            modification_map = {}
            for mod in modifications:
                modification_map[mod['original_text']] = mod['new_text']
            
            print(f"开始复制文档内容，修改条目数: {len(modifications)}")
            
            # 首先复制文档的核心属性和样式
            self._copy_document_properties()
            
            # 遍历原文档的所有元素（段落、表格、图片等）
            element_count = 0
            for element in self.original_doc.element.body:
                element_count += 1
                element_tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
                
                if element.tag.endswith('p'):  # 段落
                    self._copy_paragraph_with_modifications(element, modification_map)
                elif element.tag.endswith('tbl'):  # 表格
                    self._copy_table_with_modifications(element, modification_map)
                elif element.tag.endswith('sectPr'):  # 节属性
                    self._copy_section_properties(element)
                else:
                    # 其他元素（如图片、图表、分页符等）直接复制
                    self._copy_element_directly(element)
                    print(f"直接复制元素类型: {element_tag}")
            
            print(f"文档复制完成，共处理元素: {element_count}")
            
        except Exception as e:
            print(f"复制和修改内容时出错: {str(e)}")
            import traceback
            traceback.print_exc()
    
    def _copy_paragraph_with_modifications(self, para_element, modification_map):
        """复制段落并应用修改，保持所有格式和内嵌对象"""
        try:
            # 获取原始段落对象
            from docx.text.paragraph import Paragraph
            original_para = Paragraph(para_element, self.original_doc)
            
            # 创建新段落
            new_para = self.modified_doc.add_paragraph()
            
            # 复制段落级别的格式
            try:
                if original_para.style:
                    new_para.style = original_para.style.name
                new_para.alignment = original_para.alignment
                
                # 复制段落格式
                if original_para.paragraph_format:
                    pf = new_para.paragraph_format
                    orig_pf = original_para.paragraph_format
                    
                    if orig_pf.space_before:
                        pf.space_before = orig_pf.space_before
                    if orig_pf.space_after:
                        pf.space_after = orig_pf.space_after
                    if orig_pf.line_spacing:
                        pf.line_spacing = orig_pf.line_spacing
                    if orig_pf.left_indent:
                        pf.left_indent = orig_pf.left_indent
                    if orig_pf.right_indent:
                        pf.right_indent = orig_pf.right_indent
                        
            except Exception as e:
                print(f"复制段落格式时出错: {str(e)}")
            
            # 处理段落中的runs和嵌入对象
            for run in original_para.runs:
                text = run.text
                
                # 应用文本修改
                modified = False
                for original_text, new_text in modification_map.items():
                    if original_text in text:
                        text = text.replace(original_text, new_text)
                        modified = True
                
                # 创建新run
                new_run = new_para.add_run(text)
                
                # 复制run格式
                self._copy_run_format(run, new_run)
                
                # 如果文本被修改，添加高亮
                if modified:
                    try:
                        new_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    except:
                        pass
                
                # 复制run中的嵌入对象（如图片）
                self._copy_run_embedded_objects(run, new_run)
                        
        except Exception as e:
            print(f"复制段落时出错: {str(e)}")
    
    def _copy_run_format(self, original_run, new_run):
        """复制run的格式"""
        try:
            # 复制字体格式
            new_run.bold = original_run.bold
            new_run.italic = original_run.italic
            new_run.underline = original_run.underline
            
            if original_run.font.name:
                new_run.font.name = original_run.font.name
            if original_run.font.size:
                new_run.font.size = original_run.font.size
            if original_run.font.color and original_run.font.color.rgb:
                new_run.font.color.rgb = original_run.font.color.rgb
            if hasattr(original_run.font, 'highlight_color') and original_run.font.highlight_color:
                new_run.font.highlight_color = original_run.font.highlight_color
                
        except Exception as e:
            print(f"复制run格式时出错: {str(e)}")
    
    def _copy_run_embedded_objects(self, original_run, new_run):
        """复制run中的嵌入对象（如图片）"""
        try:
            # 检查原始run中是否有嵌入对象
            run_xml = original_run._element.xml
            
            # 如果包含图片或其他对象，尝试复制
            if any(keyword in run_xml for keyword in ['drawing', 'pict', 'object', 'blip', 'pic:']):
                print("发现嵌入对象，开始复制...")
                
                from copy import deepcopy
                from lxml import etree
                
                # 复制原始run的所有子元素
                for child in original_run._element:
                    child_tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                    
                    if any(tag in child.tag for tag in ['drawing', 'pict', 'object']):
                        try:
                            # 深度复制嵌入对象
                            copied_child = deepcopy(child)
                            new_run._element.append(copied_child)
                            print(f"成功复制嵌入对象: {child_tag}")
                        except Exception as e:
                            print(f"复制嵌入对象 {child_tag} 失败: {str(e)}")
                            # 尝试直接复制
                            try:
                                new_run._element.append(child)
                                print(f"使用直接复制成功: {child_tag}")
                            except Exception as e2:
                                print(f"直接复制也失败: {str(e2)}")
                
                # 额外检查：如果run中有图片引用，确保资源也被复制
                if 'blip' in run_xml or 'pic:' in run_xml:
                    self._copy_image_resources(original_run, new_run)
                        
        except Exception as e:
            print(f"复制嵌入对象时出错: {str(e)}")
    
    def _copy_image_resources(self, original_run, new_run):
        """复制图片资源"""
        try:
            # 获取原文档中的图片资源
            if hasattr(self.original_doc, 'part') and hasattr(self.modified_doc, 'part'):
                orig_part = self.original_doc.part
                mod_part = self.modified_doc.part
                
                # 复制相关的图片资源
                for rel in orig_part.rels.values():
                    if hasattr(rel, 'target_part') and hasattr(rel.target_part, '_blob'):
                        # 这是一个图片或其他媒体资源
                        try:
                            # 添加资源到新文档
                            mod_part.rels.add_relationship(rel.reltype, rel.target_part, rel.rId)
                            print(f"复制资源关系: {rel.rId}")
                        except Exception as e:
                            print(f"复制资源关系失败: {str(e)}")
                            
        except Exception as e:
            print(f"复制图片资源时出错: {str(e)}")
    
    def _copy_table_with_modifications(self, table_element, modification_map):
        """复制表格并应用修改"""
        try:
            from docx.table import Table
            original_table = Table(table_element, self.modified_doc)
            
            # 创建新表格
            new_table = self.modified_doc.add_table(
                rows=len(original_table.rows), 
                cols=len(original_table.columns)
            )
            
            # 复制表格样式
            try:
                if original_table.style:
                    new_table.style = original_table.style.name
            except:
                pass
            
            # 复制表格内容
            for row_idx, row in enumerate(original_table.rows):
                for col_idx, cell in enumerate(row.cells):
                    new_cell = new_table.cell(row_idx, col_idx)
                    
                    # 清空新单元格的默认内容
                    new_cell._element.clear_content()
                    
                    # 复制单元格的所有段落
                    for para in cell.paragraphs:
                        new_para = new_cell.add_paragraph()
                        
                        # 复制段落样式
                        try:
                            if para.style:
                                new_para.style = para.style.name
                            new_para.alignment = para.alignment
                        except:
                            pass
                        
                        # 复制段落内容
                        for run in para.runs:
                            text = run.text
                            
                            # 应用修改
                            modified = False
                            for original_text, new_text in modification_map.items():
                                if original_text in text:
                                    text = text.replace(original_text, new_text)
                                    modified = True
                            
                            new_run = new_para.add_run(text)
                            
                            # 复制格式
                            self._copy_run_format(run, new_run)
                            
                            # 高亮修改内容
                            if modified:
                                try:
                                    new_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                                except:
                                    pass
                            
                            # 复制嵌入对象
                            self._copy_run_embedded_objects(run, new_run)
                            
        except Exception as e:
            print(f"复制表格时出错: {str(e)}")
    
    def _copy_document_properties(self):
        """复制文档属性和设置"""
        try:
            # 复制文档核心属性
            if hasattr(self.original_doc, 'core_properties') and hasattr(self.modified_doc, 'core_properties'):
                orig_props = self.original_doc.core_properties
                mod_props = self.modified_doc.core_properties
                
                if orig_props.title:
                    mod_props.title = orig_props.title
                if orig_props.author:
                    mod_props.author = orig_props.author
                if orig_props.subject:
                    mod_props.subject = orig_props.subject
                if orig_props.comments:
                    mod_props.comments = orig_props.comments
                    
        except Exception as e:
            print(f"复制文档属性时出错: {str(e)}")
    
    def _copy_section_properties(self, sectPr_element):
        """复制节属性（页面设置等）"""
        try:
            # 复制节属性到新文档
            from copy import deepcopy
            new_sectPr = deepcopy(sectPr_element)
            self.modified_doc.element.body.append(new_sectPr)
            print("已复制节属性（页面设置）")
        except Exception as e:
            print(f"复制节属性时出错: {str(e)}")
    
    def _copy_element_directly(self, element):
        """直接复制文档元素（如图片、图表等）"""
        try:
            from copy import deepcopy
            from lxml import etree
            
            # 获取元素类型用于调试
            element_tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
            
            # 深度复制元素以避免引用问题
            new_element = deepcopy(element)
            
            # 将复制的元素添加到新文档中
            self.modified_doc.element.body.append(new_element)
            
            print(f"成功复制元素: {element_tag}")
            
        except Exception as e:
            print(f"直接复制元素时出错: {str(e)}")
            # 尝试简单复制
            try:
                self.modified_doc.element.body.append(element)
                print("使用简单复制方法成功")
            except Exception as e2:
                print(f"简单复制也失败: {str(e2)}")
    
    def save_modified_document(self, output_path: str) -> Tuple[bool, str]:
        """保存修改后的文档"""
        try:
            if not self.modified_doc:
                return False, get_text('no_document')
            
            self.modified_doc.save(output_path)
            return True, get_text('download_started')
            
        except Exception as e:
            return False, f"{get_text('processing_failed')}: {str(e)}"
    
    def get_document_info(self) -> Dict[str, Any]:
        """获取文档信息"""
        if not self.original_doc:
            return {}
        
        return {
            'paragraph_count': len(self.original_doc.paragraphs),
            'table_count': len(self.original_doc.tables),
            'image_count': len(self.images),
            'style_count': len(self.styles),
            'modifications_count': len(self.modifications)
        }
    
    def cleanup_temp_files(self):
        """清理临时文件"""
        try:
            # 清理图片缓存
            self.images.clear()
            
            # 清理其他临时数据
            self.tables.clear()
            self.styles.clear()
            
        except Exception as e:
            print(f"清理临时文件时出错: {str(e)}") 